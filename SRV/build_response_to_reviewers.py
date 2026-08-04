#!/usr/bin/env python3
"""Build the point-by-point response for manuscript EST-D-26-05562."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "Response_to_Reviewers.docx"
TITLE = (
    "An Evidence-Gated Assessment of Proposed Twisted Bilayer Graphene "
    "Coatings for Hard-Carbon Sodium-Ion Anodes"
)


def font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_text(doc, text, *, bold=False, italic=False, size=11, color=None, space_after=5):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    font(run, size=size, bold=bold, italic=italic, color=color)
    return paragraph


def add_response(doc, number, comment, response, locations):
    heading = doc.add_paragraph(style="Response Comment")
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(f"Comment {number}")
    font(run, size=11, bold=True, color=(31, 78, 121))

    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Inches(0.28)
    quote.paragraph_format.right_indent = Inches(0.15)
    quote.paragraph_format.space_after = Pt(4)
    quote_run = quote.add_run(comment)
    font(quote_run, size=10.5, italic=True, color=(80, 80, 80))

    answer = doc.add_paragraph()
    answer.paragraph_format.space_after = Pt(3)
    answer.paragraph_format.line_spacing = 1.08
    label = answer.add_run("Response: ")
    font(label, bold=True, color=(0, 100, 80))
    body = answer.add_run(response)
    font(body)

    change = doc.add_paragraph()
    change.paragraph_format.space_after = Pt(8)
    label = change.add_run("Changes in the revised manuscript: ")
    font(label, bold=True)
    body = change.add_run(locations)
    font(body)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(11)
    if "Response Comment" not in styles:
        response_style = styles.add_style("Response Comment", WD_STYLE_TYPE.PARAGRAPH)
        response_style.base_style = styles["Heading 3"]
        response_style.paragraph_format.space_before = Pt(8)
        response_style.paragraph_format.space_after = Pt(3)

    title = add_text(doc, "Response to the Editor and Reviewers", bold=True, size=16, space_after=2)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = add_text(doc, "Manuscript EST-D-26-05562", bold=True, size=12, space_after=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    manuscript_title = add_text(doc, TITLE, italic=True, size=11, space_after=10)
    manuscript_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_text(
        doc,
        "Dear Professor Cabeza and Reviewers,",
        space_after=5,
    )
    add_text(
        doc,
        "Thank you for the detailed and constructive reviews. We agreed that the previous manuscript connected a planar aqueous model-system observation to downstream battery, economic, and innovation-system questions without a validated bridge. We therefore made a fundamental methodological revision rather than add further unsupported assumptions. The paper is now an evidence-gated assessment of a proposed coating, not a claim that a tBLG-coated sodium-ion battery has been demonstrated or that a complete process TEA can be performed from the available files.",
    )
    add_text(doc, "The principal global changes are:", bold=True, space_after=3)
    for item in [
        "The title and abstract now state that the coating is proposed and unvalidated.",
        "The Yu et al. redox system was corrected; the approximately 10× statement, 5× derating, and universal kinetic multipliers were removed.",
        "A Level 0–6 evidence hierarchy now prevents propagation to battery performance, manufacturing, economics, commercialization, or policy without an independently measured mapping.",
        "The cost calculation is labelled an illustrative parametric stress test; the 2.43× lifetime-throughput and other non-harmonized break-even arithmetic were deleted.",
        "Numerical TIS scores were removed from the results and replaced by a provenance/evidence-availability audit.",
        "Table 1, Figure 5, the graphical abstract, references, Data availability statement, AI declaration, highlights, and cover letter were rebuilt for consistency.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        font(p.add_run(item))

    add_text(
        doc,
        "We use section and table identifiers below because continuous line numbers have been added and may repaginate when the editorial system creates its PDF.",
        italic=True,
        size=10.5,
        color=(80, 80, 80),
        space_after=10,
    )

    h = doc.add_paragraph(style="Heading 1")
    font(h.add_run("Response to the Editor"), size=14, bold=True, color=(31, 78, 121))
    add_response(
        doc,
        "E1",
        "Please consider all reviewer comments, describe every change point by point, and justify any comment not adopted.",
        "We have addressed every numbered and general comment below. Where reviewers requested a complete plant TEA or a quantitative performance-to-value mapping, we agree those analyses would be necessary for such claims. However, the project contains no validated tBLG-on-powder process, mass/energy balance, equipment design, or measured battery benefit. Inventing these inputs would reduce rather than improve rigor. We therefore withdrew the corresponding TEA/competitiveness claims, explicitly documented the missing evidence, and converted the work to an evidence-gated stress test.",
        "Title; Abstract; Sections 1.3, 2.1–2.4, 3.4, 4, 5, and 6; Table 1; Figure 5; Data availability; and the separate revised files.",
    )

    h = doc.add_paragraph(style="Heading 1")
    font(h.add_run("Response to Reviewer 1"), size=14, bold=True, color=(31, 78, 121))
    add_response(
        doc,
        "1-General",
        "The manuscript lacks the mass/energy balances, energy consumption, capacity, CapEx, OpEx, and economic indicators expected in a comprehensive TEA.",
        "We agree with the reviewer about what a comprehensive process TEA requires. Those inputs cannot be derived defensibly because no tBLG coating unit operation on hard-carbon powder has been demonstrated. We therefore no longer classify the calculation as a comprehensive TEA. It is now an illustrative parametric cost stress test that propagates archived accounting assumptions, with all excluded plant and application variables itemized. This is a substantive narrowing of the research claim, not an attempt to fill missing process data with conjecture.",
        "Abstract; Sections 2.2, 3.2–3.4, 5, and 6; Table A.1; revised title and cover letter.",
    )
    add_response(
        doc,
        "1",
        "Table 1 should be supported by references.",
        "Agreed. We also recognized that the prior table mixed active-material, cell, and system boundaries. It has been replaced by a Level 0–6 evidence hierarchy. Every published evidence row now carries directly relevant citations, while unsupported tBLG battery levels are explicitly marked as not met.",
        "Section 3.1 and rebuilt Table 1.",
    )
    add_response(
        doc,
        "2",
        "The TEA results are limited; discuss cost categories, drivers, reduction strategies, materials, alternative coatings, and performance–economics trade-offs.",
        "We expanded the transparent reporting of the five archived categories and retained the reproducible Monte Carlo summary and Spearman rank associations. We use 'association' rather than variance contribution or causal driver. We do not compare alternative coatings or claim performance–cost trade-offs because the archive lacks recipes, loadings, yields, and matched performance data. A directly relevant Al2O3 ALD study is cited only as evidence that conventional coating claims require a specified process and cell protocol. CVD coating is identified as the strongest association within the assumed model, not as a validated real-world cost driver.",
        "Sections 1.1, 2.2, 3.2, and 3.3; Figure 2; Table 2; Table A.1; new direct ALD reference.",
    )
    add_response(
        doc,
        "3",
        "Methods should state the software, mass/energy-balance method, plant scale, annual hours, location, application, and assumptions.",
        "The revised Methods identify Python, NumPy, pandas, the random seed, sample count, distributions, formula, functional unit, and factory-gate accounting boundary. The absence of a process route means that mass/energy balances, plant scale, operating hours, location, equipment, CapEx, and OpEx are not merely omitted details; they are currently unidentifiable. We state this explicitly and require them at Evidence Level 5 before a process TEA. No specific application is selected because no validated full-cell data exist for an application model.",
        "Sections 2.2 and 2.4; Gate 5 in Section 4.3; Table 1; Table A.1.",
    )
    add_response(
        doc,
        "4",
        "The sensitivity-analysis method is not provided.",
        "We now describe the 10,000-run Monte Carlo sampling and Spearman rank correlations between each primitive input and total scenario cost. We explicitly state that these correlations are monotonic input–output associations under the assumed distributions, not a variance decomposition or causal sensitivity analysis. Hard-coded uncertainty bands and tornado percentages are retained only as historical inputs and are not interpreted statistically.",
        "Section 2.2, Figure 2 caption, Figure 3 caption, and the reproducible source files.",
    )
    add_response(
        doc,
        "5",
        "The LCA discussion should not appear as a result because no LCA was performed.",
        "Agreed. No LCA result is reported. The manuscript now states in Limitations that a coating-specific life-cycle inventory is unavailable and cites conventional battery LCA studies only as future methodological baselines.",
        "Section 5 and Evidence Level/Gate 5; no LCA subsection remains in Results.",
    )
    add_response(
        doc,
        "6",
        "The large number of TEA assumptions should be shown and their effects explained.",
        "Agreed. Table A.1 lists every archived distribution and labels its evidence status. The Methods distinguish assumed central values from validated data, and Results report what the distributions mathematically produce without treating the output as a forecast. The principal consequence of these assumptions is now stated directly: the numerical envelope characterizes the archived model, not a realizable process.",
        "Sections 2.2, 3.2–3.3, and 5; Table A.1; supplementary CSV and source code.",
    )

    h = doc.add_paragraph(style="Heading 1")
    font(h.add_run("Response to Reviewer 2"), size=14, bold=True, color=(31, 78, 121))
    reviewer2 = [
        ("1", "The first two Abstract sentences are confusing and appear to treat twist-dependent enhancement as a hard-carbon coating.", "The Abstract was rewritten. It first states the planar aqueous model-system observation, then immediately states that no tBLG-coated hard-carbon SIB fabrication or test data were identified. It no longer treats the observation as a coating result or uses an optimistic-learning-rate qualifier as evidence.", "Abstract."),
        ("2", "The Abstract lacks a coherent logic and does not clearly present the novelty.", "The revised Abstract follows one logic: supported Level 0 evidence, unmet downstream levels, a conditional cost stress test with no performance credit, a provenance audit, and the falsification-oriented output. The novelty is the evidence-gating method rather than a claim of material performance or competitiveness.", "Title and Abstract."),
        ("3", "The discussion of existing SIB interface strategies, carbon coatings, maturity, and marginal cost is confusing.", "We removed the unsupported maturity and marginal-cost language. The revision now distinguishes surface/electrolyte engineering from a directly demonstrated Al2O3 ALD hard-carbon example, and states that any coating comparison requires specified material, loading, process, and cell conditions.", "Section 1.1 and direct ALD reference."),
        ("4", "It is unclear whether a precise twist-controlled bilayer can form on irregular hard carbon, remain exposed after SEI growth, and avoid blocking Na+ transport.", "Agreed. These are now explicit failure modes and Gate 2 measurements. The manuscript makes no manufacturability claim and requires coverage, twist distribution, defects, adhesion, wetting, pore access, processing survival, SEI behavior, and cycling stability before progression.", "Sections 1.1 and 3.5; Table 1; Gate 2; Limitations."),
        ("5", "The abrupt ALD Al2O3/TiO2 discussion is poorly connected and the TRL 5–8 claim is unclear.", "The generic Al2O3/TiO2 and TRL comparison was deleted. A single directly relevant Al2O3 ALD hard-carbon study is used only to illustrate the evidence needed for a conventional coating claim; no maturity level is assigned.", "Section 1.1 and corrected reference list."),
        ("6", "Figure 1 dates and shaded uncertainty regions are unsupported.", "Agreed. Figure 1 contains no dates, TRLs, or statistical bands. It now uses a non-sequential evidence-boundary schematic to separate the supported planar model system from untested particle, cell, process, commercial, and policy domains.", "Figure 1 and caption; Section 3.1."),
        ("7", "Table 1 compares capacities at inconsistent material/cell levels.", "Agreed. The performance comparison was removed entirely. Table 1 is now an evidence hierarchy, so active-material, cathode, cell, and system values are no longer placed in one comparison.", "Rebuilt Table 1 and Section 3.1."),
        ("8", "The TIS analysis does not support its claimed purpose of assessing industrialization.", "Agreed. The TIS component is no longer presented as an industrialization assessment. It is a provenance audit that identifies missing records and the conditions required for any future empirical TIS study.", "Sections 2.3 and 4; Figure 5; Table A.2."),
        ("9", "'Beyond prototype' is undefined and the manuscript lacks a central argument.", "We define proof of concept, prototype, and commercial relevance in Section 1.3. The central argument is now explicit: downstream inference is prohibited unless an independently measured mapping connects adjacent evidence levels.", "Sections 1.3 and 2.1; Table 1."),
        ("10", "The USD/kWh system boundary is undefined.", "The functional unit and boundary are now defined as 1 kWh of nominal battery-system capacity at the factory gate, with the five included accounting categories and every excluded installation, financing, operation, maintenance, replacement, end-of-life, and process-design item listed. We emphasize that the BOS label comes from the archive and that the result is not a levelized cost or vendor quote.", "Section 2.2 and Table A.1."),
        ("11", "The analysis does not quantify the economic value of performance improvements in a unified system model.", "We agree that such a model would be required for a competitiveness claim. Because no validated tBLG full-cell performance exists, we did not create a speculative benefit model. Instead, we deleted the 2.43× lifetime-throughput and LFP ratio calculations, assign zero performance credit, and state the full-cell/application data required before value can be quantified. The manuscript no longer concludes whether the concept is competitive.", "Sections 2.2, 3.4, 5, and 6; prior break-even arithmetic removed."),
        ("12", "The aqueous outer-sphere electron-transfer result cannot be mapped to coupled hard-carbon SIB processes.", "Agreed. On checking the primary paper, we also found that our earlier Fe(CN)6 description was inaccurate: Yu et al. used aqueous solution-phase probes including Ru(NH3)6³⁺/²⁺ and reported probe-dependent behavior. The corrected manuscript treats this only as Level 0 evidence and introduces no transfer coefficient to desolvation, SEI, adsorption, insertion, pore filling, capacity, rate, or life. The older synthetic project files were not used as validation.", "Sections 1, 2.1, 3.1, 5, and 6; Table 1; corrected Yu et al. citation."),
        ("13", "Claims of 'first integration' and quantitative prediction lack a systematic search; policy prescriptions are premature.", "All priority claims were removed. Funding levels, named institutional leadership, standards dates, and procurement recommendations are absent. Gate 6 explicitly requires traceable indicators, independent coding, and stakeholder evidence before policy analysis.", "Sections 1.3, 2.4, 4, 5, and 6."),
        ("14", "The cost categories do not correspond to an actual production process.", "Agreed. We no longer describe them as a production-process model. They are archived accounting proxies in a parametric stress test. A defined unit operation, process flow, balances, yield, equipment, CapEx/OpEx, and quality control are required at Gate 5 before process relevance is claimed.", "Sections 2.2, 3.4–3.5, 4.3, and 5."),
    ]
    for number, comment, response, locations in reviewer2:
        add_response(doc, number, comment, response, locations)

    h = doc.add_paragraph(style="Heading 1")
    font(h.add_run("Response to Reviewer 3 (attached report)"), size=14, bold=True, color=(31, 78, 121))
    reviewer3 = [
        ("1", "The Abstract should state clearly that no experimental validation of a tBLG-coated hard-carbon SIB exists.", "Agreed. This is now stated in the first sentence context of the Abstract, and the closing sentence enumerates the conclusions that are not supported.", "Abstract and Conclusion."),
        ("2", "The title may overstate the maturity and competitiveness scope.", "Agreed. The new title uses 'Evidence-Gated Assessment' and 'Proposed ... Coatings' and does not ask whether the technology is competitive.", "Title, graphical abstract, highlights, cover letter, and data-package README."),
        ("3", "The aqueous-to-battery assumption and 5× derating need greater caution; test 1×, 1.5×, and 2× cases.", "We agree with the concern and went further than adding more arbitrary sensitivity cases. The 5× derating and universal 1×/1.5×/2× multipliers were removed because different endpoints have different units and favorable directions. Future tests instead use endpoint-specific effect sizes such as log impedance/diffusion ratios and absolute efficiency or retention differences, with minimum important differences and confidence intervals defined before testing.", "Section 2.1, Section 3.1, and Limitations."),
        ("4", "Manufacturability on hard-carbon particles needs deeper treatment.", "Agreed. The revised discussion covers coating coverage, twist distribution, defects, adhesion, transfer residue, powder/slurry processing, calendaring, wetting, pore blocking, SEI evolution, Na+ transport, inactive mass, and cycling stability. These are Gate 2 failure tests, not assumed capabilities.", "Sections 1.1 and 3.5; Table 1; Gate 2; Limitations."),
        ("5", "The Monte Carlo distributions and especially CVD cost need more transparent justification.", "The exact distributions, central values, evidence status, random seed, run count, formula, and rank-association method are now reported. We explicitly call the CVD distribution an unvalidated scale-up assumption because no tBLG-powder process or quote exists. Its numerical influence is reported only conditional on that assumption.", "Section 2.2; Table A.1; Figures 2–3; supplementary code and CSV files."),
        ("6", "The 1–7 TIS scoring, raw indicators, weights, and lack of expert elicitation require stronger support.", "We agree that the archived material cannot support quantitative scores. Rather than defend the decimal ratings, we removed them from the manuscript's results and replaced them with an evidence-availability/provenance audit. Raw indicators, codebook, weights, independent coding, and stakeholder elicitation are required before comparative or policy inference.", "Sections 2.3 and 4; Figure 5; Table A.2."),
        ("7", "The graphical abstract overemphasizes a 10× charge-transfer boost.", "Agreed. The graphical abstract contains no 10× statement. It shows only Level 0 planar aqueous evidence, the missing transfer function, and the evidence gates required before battery or downstream claims.", "Revised graphical-abstract PowerPoint file and matching PNG/PDF exports."),
    ]
    for number, comment, response, locations in reviewer3:
        add_response(doc, number, comment, response, locations)

    h = doc.add_paragraph(style="Heading 1")
    font(h.add_run("Additional integrity and submission corrections"), size=14, bold=True, color=(31, 78, 121))
    add_text(
        doc,
        "During the citation audit we found several reference entries whose DOI metadata pointed to unrelated papers or could not be verified. These entries were replaced with directly relevant primary or review sources; the Vaalma DOI was corrected; author lists were standardized; and all 33 references were renumbered by first appearance across body text and tables. The public Zenodo v1.0.1 record was also checked and found to predate the revised package, so the Data availability statement no longer claims that the current public version contains the new files. The unchanged concept DOI is retained, and publication of a new version is listed as a required pre-submission action.",
        space_after=8,
    )
    add_text(
        doc,
        "We hope that the reconstruction resolves the reviewers' central concern: the revised paper no longer uses a planar aqueous model observation as if it were evidence of a sodium-ion battery, industrial process, competitive product, or policy case.",
        space_after=10,
    )
    add_text(doc, "Sincerely,", space_after=2)
    add_text(doc, "Qinyou Yang, on behalf of all authors", space_after=0)

    doc.core_properties.title = "Response to reviewers for EST-D-26-05562"
    doc.core_properties.subject = TITLE
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
    print(f"Wrote {OUTPUT}")
