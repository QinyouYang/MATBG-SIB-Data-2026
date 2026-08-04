#!/usr/bin/env python3
"""Apply final layout-only fixes to the already revised manuscript.

This script is intentionally idempotent. It does not rebuild scientific text,
citations, references, or tables; it only keeps the nomenclature table together
and places the unchanged Zenodo concept DOI on its own line.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
MANUSCRIPT = ROOT / "Manuscript.docx"

DATA_TEXT = (
    "The screening inputs, deterministic source code, recomputed summaries, and "
    "evidence-provenance table used in this revision are provided with the "
    "submission as Supplementary Material. They are modeled inputs and outputs, "
    "not measured tBLG-SIB data. Zenodo v1.0.1 [33] is the currently public "
    "version under the unchanged concept DOI:"
)
DOI_TEXT = "https://doi.org/10.5281/zenodo.19362805"
CONTINUATION_TEXT = (
    "This public version predates the present revision and must not be represented "
    "as containing the revised supplementary package. The revised package is "
    "prepared for release as a new version under the same concept DOI before "
    "resubmission."
)


def paragraph_with_prefix(doc, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(prefix):
            return paragraph
    raise ValueError(f"Could not find paragraph starting with {prefix!r}")


def insert_before(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    result = Paragraph(new_p, paragraph._parent)
    result.style = "Normal"
    result.add_run(text)
    return result


def ensure_nomenclature_page_break(doc) -> None:
    table = doc.tables[0]
    previous = table._tbl.getprevious()
    if previous is not None and previous.xpath(".//w:br[@w:type='page']"):
        return
    new_p = OxmlElement("w:p")
    new_r = OxmlElement("w:r")
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    new_r.append(page_break)
    new_p.append(new_r)
    table._tbl.addprevious(new_p)


def split_data_availability(doc) -> None:
    data_heading = paragraph_with_prefix(doc, "Data availability")
    ai_heading = paragraph_with_prefix(
        doc,
        "Declaration of generative AI and AI-assisted technologies",
    )
    data_body = Paragraph(data_heading._p.getnext(), data_heading._parent)
    data_body.clear()
    data_body.add_run(DATA_TEXT)
    data_body.style = "Normal"

    current = data_body._p.getnext()
    while current is not None and current is not ai_heading._p:
        following = current.getnext()
        current.getparent().remove(current)
        current = following

    doi_paragraph = insert_before(ai_heading, DOI_TEXT)
    continuation = insert_before(ai_heading, CONTINUATION_TEXT)
    for paragraph in (data_body, doi_paragraph, continuation):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Inches(0)
    doi_paragraph.paragraph_format.keep_with_next = True
    doi_run = doi_paragraph.runs[0]
    doi_run.font.name = "Times New Roman"
    doi_run.font.size = Pt(12)
    doi_fonts = doi_run._element.get_or_add_rPr().get_or_add_rFonts()
    doi_fonts.set(qn("w:ascii"), "Times New Roman")
    doi_fonts.set(qn("w:hAnsi"), "Times New Roman")


def main() -> None:
    doc = Document(MANUSCRIPT)
    ensure_nomenclature_page_break(doc)
    split_data_availability(doc)
    doc.save(MANUSCRIPT)
    print("Applied final manuscript layout fixes.")


if __name__ == "__main__":
    main()
