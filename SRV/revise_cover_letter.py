"""Revise the Journal of Energy Storage cover letter to match the final manuscript."""

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
TARGET = ROOT / "Cover Letter.docx"


def set_font(run, name="Times New Roman", size=11, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph(paragraph, text, *, bold=False, align=None):
    paragraph.clear()
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_font(run, size=12 if bold else 11, bold=bold)


def main():
    doc = Document(TARGET)
    for section in doc.sections:
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.82)
        section.right_margin = Inches(0.82)
    text = [
        ("Cover letter", True, WD_ALIGN_PARAGRAPH.CENTER),
        ("Qinyou Yang", False, None),
        ("School of Physics and Electronic Engineering, Mudanjiang Normal University", False, None),
        ("No. 191 Wenhua Street, Aimin District, Mudanjiang 157011, Heilongjiang, China", False, None),
        ("July 31, 2026", False, None),
        ("Re: Major revision of Manuscript EST-D-26-05562", True, None),
        ("Dear Professor Cabeza,", False, None),
        (
            "We are pleased to resubmit the substantially revised manuscript entitled “An Evidence-Gated Assessment of Proposed Twisted Bilayer Graphene Coatings for Hard-Carbon Sodium-Ion Anodes” for consideration by Journal of Energy Storage. We thank you and the reviewers for identifying the need to reconstruct the evidence chain. The manuscript is not published or under consideration elsewhere, and all authors approve this revision.",
            False,
            None,
        ),
        (
            "The revision treats the published planar-tBLG aqueous outer-sphere electron-transfer observation strictly as Level 0 model-system evidence. We corrected the redox-probe description; removed the approximately 10× statement, 5× derating assumption, and universal kinetic multipliers; and introduce no mapping to Na+ desolvation, SEI transport, impedance, capacity, rate capability, or cycle life. No tBLG-coated hard-carbon sodium-ion-battery fabrication or test data were identified.",
            False,
            None,
        ),
        (
            "We reclassified the archived 10,000-run calculation as an illustrative parametric cost stress test, not a bottom-up plant TEA or commercial forecast. Because no validated coating process exists from which to derive mass/energy balances, equipment sizing, CapEx, or OpEx, these are reported as unmet evidence requirements. The earlier cost-ratio and lifetime-throughput calculations were deleted, and no performance credit is coupled to the cost model.",
            False,
            None,
        ),
        (
            "Numerical Technology Innovation Systems scores were replaced by a provenance audit because the underlying indicators, codebook, weights, independent coding, and stakeholder evidence are unavailable. The revised contribution is a Level 0–6 hierarchy with explicit falsification and progression gates from relevant-electrolyte replication through particle coating, cell validation, process TEA/LCA, and stakeholder-based analysis. It supports no commercialization ranking or policy recommendation.",
            False,
            None,
        ),
        (
            "We verified every in-text citation against the reference list, corrected erroneous DOI metadata, added directly relevant recent sources, and renumbered all 33 references by first appearance. A point-by-point response and reproducible supplementary package accompany the revision. As requested, the Zenodo concept DOI remains unchanged: https://doi.org/10.5281/zenodo.19362805. The revised package is prepared for publication as a new version under that concept record. The manuscript also includes the required generative-AI disclosure.",
            False,
            None,
        ),
        (
            "We believe the revised article fits Journal of Energy Storage as a transparent, boundary-aware assessment of an emerging storage-material proposal. It identifies the measurements and process records needed before an interfacial observation can support battery, economic, or innovation-system inference. This research was supported by the Key Project of Heilongjiang Provincial Natural Science Foundation United Fund (Grant No. ZL2025C004), and the authors declare no competing interests.",
            False,
            None,
        ),
        ("Please address all correspondence to the corresponding author, Qinyou Yang, at qinyou_yang@163.com.", False, None),
        ("Thank you for your consideration.", False, None),
        ("Sincerely,", False, None),
        ("Qinyou Yang", False, None),
        ("On behalf of all co-authors", False, None),
    ]
    while len(doc.paragraphs) < len(text):
        doc.add_paragraph()
    for paragraph, (content, bold, align) in zip(doc.paragraphs, text):
        set_paragraph(paragraph, content, bold=bold, align=align)
    for paragraph in doc.paragraphs[len(text):]:
        paragraph._element.getparent().remove(paragraph._element)
    doc.save(TARGET)


if __name__ == "__main__":
    main()
