#!/usr/bin/env python3
"""Apply the evidence-gated Journal of Energy Storage revision.

The script edits the submission copy in place, but it never creates experimental
evidence or publishes a repository release. It corrects the model-system redox
description, removes unsupported battery/commercial/policy mappings, replaces
the numerical TIS rubric with a provenance audit, rebuilds the reference list,
renumbers citations in true document order, applies three-line table styling,
and refreshes the Highlights file.
"""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
MANUSCRIPT = ROOT / "Manuscript.docx"
HIGHLIGHTS = ROOT / "Highlights.docx"
SUBMISSION_FIGURES = ROOT / "Submission_Figures"

TITLE = (
    "An Evidence-Gated Assessment of Proposed Twisted Bilayer Graphene "
    "Coatings for Hard-Carbon Sodium-Ion Anodes"
)


def set_text(paragraph: Paragraph, text: str, style: str | None = None) -> None:
    paragraph.clear()
    paragraph.add_run(text)
    if style is not None:
        paragraph.style = style


def find_paragraph(doc: _Document, *prefixes: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if any(paragraph.text.startswith(prefix) for prefix in prefixes):
            return paragraph
    raise ValueError(f"Could not find paragraph starting with: {prefixes}")


def replace_matching(
    doc: _Document,
    prefixes: tuple[str, ...],
    text: str,
    style: str | None = None,
) -> Paragraph:
    paragraph = find_paragraph(doc, *prefixes)
    set_text(paragraph, text, style)
    return paragraph


def insert_paragraph_before(paragraph: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_paragraph.style = style
    new_paragraph.add_run(text)
    return new_paragraph


def insert_page_break_before_table(table: Table) -> None:
    """Force a compact table to start on a fresh page in Word and LibreOffice."""
    new_p = OxmlElement("w:p")
    new_r = OxmlElement("w:r")
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    new_r.append(page_break)
    new_p.append(new_r)
    table._tbl.addprevious(new_p)


def next_paragraph(paragraph: Paragraph) -> Paragraph:
    sibling = paragraph._p.getnext()
    while sibling is not None and not isinstance(sibling, CT_P):
        sibling = sibling.getnext()
    if sibling is None:
        raise ValueError(f"Paragraph has no following paragraph: {paragraph.text}")
    return Paragraph(sibling, paragraph._parent)


def ensure_before(
    doc: _Document,
    heading_prefix: str,
    marker_prefix: str,
    text: str,
    style: str = "Normal",
) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(marker_prefix):
            set_text(paragraph, text, style)
            return paragraph
    return insert_paragraph_before(find_paragraph(doc, heading_prefix), text, style)


def set_author_line(paragraph: Paragraph) -> None:
    paragraph.clear()
    authors = ["Qinyou Yang", "Xinran Wang", "Niwei Zhu", "Wulin Zhang"]
    for index, author in enumerate(authors):
        if index:
            paragraph.add_run(", ")
        paragraph.add_run(author)
        affiliation = paragraph.add_run("a,*" if index == 0 else "a")
        affiliation.font.superscript = True


def set_affiliation(paragraph: Paragraph) -> None:
    paragraph.clear()
    marker = paragraph.add_run("a")
    marker.font.superscript = True
    paragraph.add_run(
        " School of Physics and Electronic Engineering, Mudanjiang Normal "
        "University, No. 191 Wenhua Street, Aimin District, Mudanjiang 157011, "
        "Heilongjiang, China"
    )


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, attrs in edges.items():
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        for key, value in attrs.items():
            element.set(qn(f"w:{key}"), str(value))


def as_three_line_table(table: Table) -> None:
    """Remove shading/vertical rules and retain only three horizontal rules."""
    try:
        table.style = "Table Normal"
    except KeyError:
        pass
    table.autofit = False
    clear = {"val": "nil", "sz": "0", "color": "FFFFFF"}
    rule = {"val": "single", "sz": "8", "color": "000000"}
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for shading in tc_pr.findall(qn("w:shd")):
                tc_pr.remove(shading)
            set_cell_border(
                cell,
                top=clear,
                bottom=clear,
                left=clear,
                right=clear,
                insideH=clear,
                insideV=clear,
            )

    title_row = table.rows[0].cells[0].text.strip().lower() == "nomenclature"
    header_index = 1 if title_row and len(table.rows) > 1 else 0
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=rule)
    for cell in table.rows[header_index].cells:
        set_cell_border(cell, bottom=rule)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
    if title_row:
        for cell in table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=rule)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def set_cell_margins(cell, *, top=35, left=60, bottom=35, right=60) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_column_widths(table: Table, widths: list[float]) -> None:
    if len(widths) != len(table.columns):
        raise ValueError("Column-width count does not match the table.")
    table.autofit = False
    for index, width_inches in enumerate(widths):
        width = Inches(width_inches)
        table.columns[index].width = width
        for cell in table.columns[index].cells:
            cell.width = width
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(width.twips))
            tc_w.set(qn("w:type"), "dxa")


def format_table_text(table: Table, size: float) -> None:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(size)
                    r_pr = run._element.get_or_add_rPr()
                    fonts = r_pr.get_or_add_rFonts()
                    fonts.set(qn("w:ascii"), "Times New Roman")
                    fonts.set(qn("w:hAnsi"), "Times New Roman")


def format_reference_list(doc: _Document) -> None:
    for number, paragraph in enumerate(reference_paragraphs(doc), 1):
        body = paragraph.text
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is not None:
            p_pr.remove(num_pr)
        set_text(paragraph, f"[{number}]  {body}", "Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.48)
        paragraph.paragraph_format.first_line_indent = Inches(-0.48)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


def remove_last_rendered_page_breaks(doc: _Document) -> None:
    for marker in list(doc.element.body.iter(qn("w:lastRenderedPageBreak"))):
        parent = marker.getparent()
        if parent is not None:
            parent.remove(marker)


def remove_row(row) -> None:
    row._tr.getparent().remove(row._tr)


def set_table_data(table: Table, rows: list[list[str]]) -> None:
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        remove_row(table.rows[-1])
    if any(len(row) != len(table.columns) for row in rows):
        raise ValueError("Table data do not match the existing column count.")
    for row, values in zip(table.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.text = value


def iter_block_items(parent: _Document):
    """Yield paragraphs and tables in their actual Word XML order."""
    if not isinstance(parent, _Document):
        raise TypeError("Only document-level traversal is supported.")
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def citation_numbers(text: str):
    for payload in re.findall(r"\[([0-9\s,;\-–]+)\]", text):
        for token in re.split(r"[,;]", payload):
            token = token.strip().replace("–", "-")
            if re.fullmatch(r"\d+(?:-\d+)?", token):
                if "-" in token:
                    low, high = map(int, token.split("-"))
                    yield from range(low, high + 1)
                else:
                    yield int(token)


def compact_citation(numbers: list[int]) -> str:
    numbers = sorted(set(numbers))
    if not numbers:
        raise ValueError("Cannot format an empty citation.")
    chunks: list[tuple[int, int]] = []
    start = end = numbers[0]
    for number in numbers[1:]:
        if number == end + 1:
            end = number
        else:
            chunks.append((start, end))
            start = end = number
    chunks.append((start, end))
    rendered: list[str] = []
    for start, end in chunks:
        if end - start >= 2:
            rendered.append(f"{start}–{end}")
        elif end == start:
            rendered.append(str(start))
        else:
            rendered.extend([str(start), str(end)])
    return "[" + ",".join(rendered) + "]"


def reference_paragraphs(doc: _Document) -> list[Paragraph]:
    ref_heading = next(i for i, p in enumerate(doc.paragraphs) if p.text == "References")
    appendix_heading = next(i for i, p in enumerate(doc.paragraphs) if p.text == "Appendix")
    return doc.paragraphs[ref_heading + 1 : appendix_heading]


def set_reference_list(doc: _Document, references: list[str]) -> None:
    existing = reference_paragraphs(doc)
    if not existing:
        raise ValueError("No reference paragraph was available as a formatting template.")
    template = deepcopy(existing[0]._p)
    appendix = find_paragraph(doc, "Appendix")
    for paragraph in existing:
        paragraph._element.getparent().remove(paragraph._element)
    for reference in references:
        new_p = deepcopy(template)
        appendix._p.addprevious(new_p)
        paragraph = Paragraph(new_p, appendix._parent)
        set_text(paragraph, reference)


def renumber_citations_and_references(doc: _Document) -> dict[int, int]:
    refs = reference_paragraphs(doc)
    reference_count = len(refs)
    ordered: list[int] = []

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph) and block.text == "References":
            break
        paragraphs = [block] if isinstance(block, Paragraph) else [
            p for row in block.rows for cell in row.cells for p in cell.paragraphs
        ]
        for paragraph in paragraphs:
            for number in citation_numbers(paragraph.text):
                if not 1 <= number <= reference_count:
                    raise ValueError(f"Citation {number} is outside 1..{reference_count}.")
                if number not in ordered:
                    ordered.append(number)

    expected = set(range(1, reference_count + 1))
    if set(ordered) != expected:
        missing = sorted(expected - set(ordered))
        raise ValueError(f"Every reference must be cited before renumbering; missing {missing}.")
    mapping = {old: new for new, old in enumerate(ordered, 1)}

    def rewrite(text: str) -> str:
        def replacement(match: re.Match) -> str:
            old_numbers = list(citation_numbers(match.group(0)))
            return compact_citation([mapping[number] for number in old_numbers])

        return re.sub(r"\[[0-9\s,;\-–]+\]", replacement, text)

    before_references = True
    for paragraph in doc.paragraphs:
        if paragraph.text == "References":
            before_references = False
        if before_references and "[" in paragraph.text:
            set_text(paragraph, rewrite(paragraph.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "[" in paragraph.text:
                        set_text(paragraph, rewrite(paragraph.text))

    texts = [paragraph.text for paragraph in refs]
    for paragraph, old_number in zip(refs, ordered):
        set_text(paragraph, texts[old_number - 1])
    return mapping


def add_continuous_line_numbers(doc: _Document) -> None:
    for section in doc.sections:
        sect_pr = section._sectPr
        for current in sect_pr.findall(qn("w:lnNumType")):
            sect_pr.remove(current)
        line_numbers = OxmlElement("w:lnNumType")
        line_numbers.set(qn("w:countBy"), "1")
        line_numbers.set(qn("w:distance"), "360")
        line_numbers.set(qn("w:restart"), "continuous")
        sect_pr.append(line_numbers)


REFERENCES = [
    "D. Larcher, J.-M. Tarascon, Towards greener and more sustainable batteries for electrical energy storage, Nat. Chem. 7 (2015) 19–29. https://doi.org/10.1038/nchem.2085.",
    "A. Yao, S.M. Benson, W.C. Chueh, Critically assessing sodium-ion technology roadmaps and scenarios for techno-economic competitiveness against lithium-ion batteries, Nat. Energy 10 (2025) 404–416. https://doi.org/10.1038/s41560-024-01701-9.",
    "International Renewable Energy Agency, Sodium-ion batteries: A technology brief, IRENA, Abu Dhabi, 2025. https://www.irena.org/Publications/2025/Nov/Sodium-ion-batteries-A-technology-brief (accessed 31 July 2026).",
    "N. Tapia-Ruiz, A.R. Armstrong, H. Alptekin, M.A. Amores, H. Au, J. Barker, et al., 2021 roadmap for sodium-ion batteries, J. Phys. Energy 3 (2021) 031503. https://doi.org/10.1088/2515-7655/ac0e32.",
    "R. Usiskin, Y. Lu, J. Popovic, M. Law, P. Balaya, Y.-S. Hu, et al., Fundamentals, status and promise of sodium-based batteries, Nat. Rev. Mater. 6 (2021) 1020–1035. https://doi.org/10.1038/s41578-021-00324-w.",
    "X. Cai, Y. Yue, Z. Yi, J. Liu, Y. Sheng, Y. Lu, Challenges and industrial perspectives on the development of sodium ion batteries, Nano Energy 129 (2024) 110052. https://doi.org/10.1016/j.nanoen.2024.110052.",
    "Z. Lu, H. Yang, Y. Guo, H. Lin, P. Shan, S. Wu, et al., Consummating ion desolvation in hard carbon anodes for reversible sodium storage, Nat. Commun. 15 (2024) 3497. https://doi.org/10.1038/s41467-024-47522-y.",
    "B. Wang, J.R. Fitzpatrick, A. Brookfield, A.J. Fielding, E. Reynolds, J. Entwistle, et al., Electron paramagnetic resonance as a tool to determine the sodium charge storage mechanism of hard carbon, Nat. Commun. 15 (2024) 3013. https://doi.org/10.1038/s41467-024-45460-3.",
    "X. Li, C. Ding, Q. Liang, J. Hu, X. Li, Y. Li, et al., Progress in hard carbons for sodium-ion batteries: Microstructure, sodium storage mechanism and initial Coulombic efficiency, J. Energy Storage 98 (2024) 112986. https://doi.org/10.1016/j.est.2024.112986.",
    "M. Liu, F. Wu, Y. Gong, Y. Li, Y. Li, X. Feng, et al., Interfacial-catalysis-enabled layered and inorganic-rich SEI on hard carbon anodes in ester electrolytes for sodium-ion batteries, Adv. Mater. 35 (2023) 2300002. https://doi.org/10.1002/adma.202300002.",
    "J.-Y. Hwang, S.-T. Myung, Y.-K. Sun, Sodium-ion batteries: present and future, Chem. Soc. Rev. 46 (2017) 3529–3614. https://doi.org/10.1039/C6CS00776G.",
    "N. Yabuuchi, K. Kubota, M. Dahbi, S. Komaba, Research development on sodium-ion batteries, Chem. Rev. 114 (2014) 11636–11682. https://doi.org/10.1021/cr500192f.",
    "D.A. Stevens, J.R. Dahn, High capacity anode materials for rechargeable sodium-ion batteries, J. Electrochem. Soc. 147 (2000) 1271–1273. https://doi.org/10.1149/1.1393348.",
    "Y. Cao, V. Fatemi, S. Fang, K. Watanabe, T. Taniguchi, E. Kaxiras, et al., Unconventional superconductivity in magic-angle graphene superlattices, Nature 556 (2018) 43–50. https://doi.org/10.1038/nature26160.",
    "J.H. Kim, S.H. Kang, D. Yoon, H. Kim, J.S. Kim, M.M. Haidari, et al., Twist angle-dependent transport properties of twisted bilayer graphene, NPG Asia Mater. 16 (2024) 36. https://doi.org/10.1038/s41427-024-00556-6.",
    "R.A. Rakkesh, P.N.B. Rebecca, T. Naveen, D. Durgalakshmi, S. Balakumar, Twisted bilayer graphene: a journey through recent advances and future perspectives, Part. Part. Syst. Charact. 41 (2024) 2300125. https://doi.org/10.1002/ppsc.202300125.",
    "L. Baldo, T. Löthman, P. Holmvall, A.M. Black-Schaffer, Defect-induced band restructuring and length scales in twisted bilayer graphene, Phys. Rev. B 108 (2023) 125141. https://doi.org/10.1103/PhysRevB.108.125141.",
    "Y. Yu, K. Zhang, H. Parks, M. Babar, S. Carr, I.M. Craig, et al., Tunable angle-dependent electrochemistry at twisted bilayer graphene with moiré flat bands, Nat. Chem. 14 (2022) 267–273. https://doi.org/10.1038/s41557-021-00865-1.",
    "L. Coello Escalante, D.T. Limmer, Microscopic origin of twist-dependent electron transfer rate in bilayer graphene, Nano Lett. 24 (2024) 14868–14874. https://doi.org/10.1021/acs.nanolett.4c04690.",
    "H. Lu, X. Chen, Y. Jia, H. Chen, Y. Wang, X. Ai, et al., Engineering Al2O3 atomic layer deposition: Enhanced hard carbon-electrolyte interface towards practical sodium ion batteries, Nano Energy 64 (2019) 103903. https://doi.org/10.1016/j.nanoen.2019.103903.",
    "W. Ma, Q. Zhang, L. Li, D. Geng, W. Hu, Small twist, big miracle—recent progress in the fabrication of twisted 2D materials, J. Mater. Chem. C 11 (2023) 15793–15816. https://doi.org/10.1039/D3TC02660D.",
    "C. Liu, Z. Li, R. Qiao, Q. Wang, Z. Zhang, F. Liu, et al., Designed growth of large bilayer graphene with arbitrary twist angles, Nat. Mater. 21 (2022) 1263–1268. https://doi.org/10.1038/s41563-022-01361-8.",
    "M.P. Hekkert, R.A.A. Suurs, S.O. Negro, S. Kuhlmann, R.E.H.M. Smits, Functions of innovation systems: a new approach for analysing technological change, Technol. Forecast. Soc. Change 74 (2007) 413–432. https://doi.org/10.1016/j.techfore.2006.03.002.",
    "A. Bergek, S. Jacobsson, B. Carlsson, S. Lindmark, A. Rickne, Analyzing the functional dynamics of technological innovation systems: a scheme of analysis, Res. Policy 37 (2008) 407–429. https://doi.org/10.1016/j.respol.2007.12.003.",
    "H. Gong, T. Hansen, The rise of China's new energy vehicle lithium-ion battery industry: The coevolution of battery technological innovation systems and policies, Environ. Innov. Soc. Transit. 46 (2023) 100689. https://doi.org/10.1016/j.eist.2022.100689.",
    "H. Gong, A.D. Andersen, The role of material resources for rapid technology diffusion in net-zero transitions: Insights from EV lithium-ion battery Technological Innovation System in China, Technol. Forecast. Soc. Change 200 (2024) 123141. https://doi.org/10.1016/j.techfore.2023.123141.",
    "C. Vaalma, D. Buchholz, M. Weil, S. Passerini, A cost and resource analysis of sodium-ion batteries, Nat. Rev. Mater. 3 (2018) 18013. https://doi.org/10.1038/natrevmats.2018.13.",
    "A. Innocenti, S. Beringer, S. Passerini, Cost and performance analysis as a valuable tool for battery material research, Nat. Rev. Mater. 9 (2024) 347–357. https://doi.org/10.1038/s41578-024-00657-2.",
    "Z. Guo, K. Zheng, M. Wang, Y. Huang, Y. Zhao, H. Au, et al., Sustainable hard carbon as anode materials for Na-ion batteries: from laboratory to upscaling, Batteries Supercaps 8 (2025) e202400428. https://doi.org/10.1002/batt.202400428.",
    "Z. Ye, H. Hijazi, W. Black, S. Azam, J.R. Dahn, M. Metzger, Impact of salts and linear carbonates on the performance of layered oxide/hard carbon sodium-ion pouch cells with alkyl carbonate electrolytes, J. Electrochem. Soc. 171 (2024) 040522. https://doi.org/10.1149/1945-7111/ad3b73.",
    "S. Wickerts, R. Arvidsson, A. Nordelöf, M. Svanström, P. Johansson, Prospective life cycle assessment of sodium-ion batteries made from abundant elements, J. Ind. Ecol. 28 (2024) 116–129. https://doi.org/10.1111/jiec.13452.",
    "J.F. Peters, M. Baumann, B. Zimmermann, J. Braun, M. Weil, The environmental impact of Li-ion batteries and the role of key parameters: a review, Renew. Sustain. Energy Rev. 67 (2017) 491–506. https://doi.org/10.1016/j.rser.2016.08.039.",
    "[software] Q. Yang, QinyouYang/MATBG-SIB-Data: MATBG-SIB Data Code, version 1.0.1, Zenodo, 2026. https://doi.org/10.5281/zenodo.19362806.",
]


def revise_manuscript() -> dict[int, int]:
    doc = Document(MANUSCRIPT)
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "Evidence-gated early-stage assessment"

    set_text(doc.paragraphs[0], TITLE)
    set_author_line(doc.paragraphs[1])
    set_affiliation(doc.paragraphs[2])
    set_text(doc.paragraphs[3], "* Corresponding author. E-mail: qinyou_yang@163.com")

    replace_matching(
        doc,
        ("Twisted bilayer graphene (tBLG) exhibits",),
        "Twisted bilayer graphene (tBLG) shows strongly angle-dependent heterogeneous electron-transfer kinetics for aqueous solution-phase outer-sphere redox probes on planar electrodes, but no source reviewed or project file contains fabrication or test data for a tBLG-coated hard-carbon sodium-ion battery. We therefore assess a proposed coating through an evidence-gated framework rather than treat it as a validated battery technology. Only the planar model-system premise is directly supported; replication in a relevant non-aqueous electrolyte, particle-level coating and structural verification, matched hard-carbon half-cell testing, practical full-cell testing, process inventory, and stakeholder evidence are successive unmet levels. Archived inputs are retained only for an illustrative factory-gate cost stress test for 1 kWh of nominal battery-system capacity, with no electrochemical-performance credit. The 10,000-run calculation gives a 2035 scenario median of 119.5 USD/kWh and a P5–P95 range of 96.5–152.6 USD/kWh; these values propagate assumed distributions and are not process-derived forecasts. No transfer coefficient is introduced from aqueous electron transfer to Na+ desolvation, solid-electrolyte-interphase transport, capacity, rate capability, or cycle life. A provenance audit further shows that raw indicators, weights, independent coding, and stakeholder elicitation needed for quantitative innovation-system or policy inference are unavailable. The defensible output is thus a falsification-oriented experimental and data-acquisition pathway, not a claim of battery performance, manufacturing feasibility, commercial competitiveness, or policy readiness.",
    )
    replace_matching(
        doc,
        ("Keywords:",),
        "Keywords: sodium-ion battery; twisted bilayer graphene; hard carbon; evidence hierarchy; parametric cost stress test; technology innovation system",
    )

    replace_matching(
        doc,
        ("The global energy transition requires",),
        "The global energy transition requires battery technologies that are cost-effective, sustainable, and resilient [1]. Sodium-ion batteries (SIBs) are being developed as complements to lithium-ion batteries for stationary storage and cost-sensitive mobility [2–6]. Hard-carbon performance is governed by coupled pore structure, surface chemistry, Na+ desolvation, electrolyte transport, charge storage, and solid-electrolyte-interphase (SEI) evolution [7–10]. These processes, rather than coating conductivity alone, determine whether an interfacial modification can benefit a practical cell.",
    )
    replace_matching(
        doc,
        ("Sodium does not form", "Graphitic layers do not"),
        "Sodium does not form the stable graphite intercalation compounds that underpin conventional graphite-based lithium-ion anodes [11–13], so twisted graphene cannot simply replace hard carbon as the active Na-storage host. tBLG has well-established twist-sensitive electronic behavior [14–17]. Yu et al. [18] reported a strong twist-angle dependence of heterogeneous electron transfer at planar tBLG electrodes, with the greatest response near the magic-angle range under their reported aqueous conditions. The study used solution-phase redox probes, including Ru(NH3)6³⁺/²⁺, and showed that the response depended on probe energy level; it did not study a battery electrolyte. Molecular modelling of twist-dependent Marcus electron transfer in bilayer graphene [19] likewise remains a model-system result. Neither study establishes Na+ desolvation, SEI transport, hard-carbon storage, or full-cell performance.",
    )
    replace_matching(
        doc,
        ("Established SIB interface strategies include",),
        "Hard-carbon interface engineering includes control of surface chemistry and porosity, electrolyte-mediated SEI design, and experimentally demonstrated coatings [7–10]. For example, an ultrathin Al2O3 layer deposited by atomic layer deposition has been studied directly on hard-carbon electrodes [20]. These examples demonstrate that a coating claim requires a specified material, loading, process, and cell protocol; they do not imply that a twist-controlled graphene bilayer is feasible or superior. The earlier generic reference to Al2O3/TiO2 as a maturity comparison is therefore withdrawn.",
    )
    replace_matching(
        doc,
        ("The tBLG concept poses",),
        "The proposed tBLG coating poses a more demanding hypothesis: a bilayer with a controlled local twist distribution would need to cover irregular hard-carbon particles, remain attached through powder processing and calendaring, permit electrolyte wetting and Na+ access, and retain any interfacial function after SEI formation. A continuous overlayer could instead block pores or add inactive mass. Existing fabrication studies address controlled tBLG primarily on comparatively flat catalyst or device substrates [15,21,22], not conformal battery-powder coatings. We therefore treat particle coating as a falsifiable manufacturing question rather than an established surface-engineering route.",
    )
    replace_matching(
        doc,
        ("Commercialization also depends",),
        "If the technical premise were later validated, commercialization would also depend on the surrounding innovation system. The Technology Innovation Systems (TIS) framework organizes seven functions concerning knowledge, networks, experimentation, direction, markets, resources, and legitimacy [23,24]. Battery-sector applications show how such functions can be studied with traceable empirical records [25,26]. TIS is not a binary industrialization test, and an unvalidated concept cannot acquire technical credibility from a favorable system score.",
    )
    replace_matching(
        doc,
        ("Prior work provides cost",),
        "Prior work provides cost, resource, industrialization, and laboratory-to-application methods for conventional SIBs and battery materials [2,6,27–29]. The unresolved question here is narrower: what evidence would be required before a proposed tBLG coating could justify process development and an application-linked assessment? We make no priority claim. Here, proof of concept means a controlled coated-hard-carbon half-cell showing a reproducible benefit in a relevant non-aqueous SIB electrolyte; a prototype requires a full cell at practical loading under a defined protocol; and commercial relevance requires reproducible process, cost, safety, environmental, and market evidence.",
    )
    replace_matching(
        doc,
        ("The central thesis is conditional",),
        "The central thesis is evidence-gated: a result at one inferential level cannot be propagated downstream without an independently measured mapping. The paper therefore separates the published planar-electrode observation from battery hypotheses, reproduces the archived cost arithmetic without assigning a performance benefit, audits the provenance of the innovation-system materials, and states decision rules that can falsify the concept. It does not calculate a battery-performance multiplier, a competitiveness threshold, or a policy roadmap.",
    )
    replace_matching(
        doc,
        ("The analysis has four parts",),
        "The analysis has four parts. Section 2.1 defines the evidence hierarchy and endpoint-specific validation rules. Section 2.2 documents the functional unit, system boundary, archived cost inputs, Monte Carlo calculation, and input-cost rank associations. Section 2.3 audits the provenance of the innovation-system materials. Section 2.4 maps electrochemical, manufacturing, full-cell, process, life-cycle, and stakeholder evidence to explicit decision gates.",
    )

    replace_matching(doc, ("2.1. Interfacial enhancement model", "2.1. Evidence hierarchy"), "2.1. Evidence hierarchy and conditional translation framework", "Heading 3")
    replace_matching(
        doc,
        ("The only direct electrochemical input",),
        "The Level 0 physical premise is the published twist-angle dependence of heterogeneous electron transfer on planar tBLG in aqueous solution-phase outer-sphere redox experiments [18], supported by a generic Marcus electron-transfer model for bilayer graphene [19]. The experimental paper interrogated Ru(NH3)6³⁺/²⁺ and other probes and found probe-dependent behavior; it did not report a sodium-ion electrolyte, a hard-carbon substrate, or a battery electrode. The earlier probe description, fixed enhancement statement, and fixed derating assumption are withdrawn because none provides a valid transfer function to the proposed application.",
    )
    replace_matching(
        doc,
        ("The scenarios are not converted",),
        "Evidence was classified by inferential distance. Level 0 is the planar aqueous-probe observation. Level 1 requires replication in a relevant non-aqueous electrolyte model system; Level 2 requires particle-level coating and structural verification; Level 3 requires matched hard-carbon half-cell testing; Level 4 requires practical full-cell testing; Level 5 requires a process inventory linked to an application model and life-cycle boundary; and Level 6 requires transparent innovation-system and stakeholder evidence. Only Level 0 is currently supported for the proposed concept. Evidence at one level is not propagated to a downstream level without an independently measured mapping.",
    )
    ensure_before(
        doc,
        "2.2.",
        "The Level 0 observation does not identify",
        "The Level 0 observation does not identify the rate-limiting step in a hard-carbon SIB electrode. Desolvation, transport through an evolving SEI, adsorption, insertion, and pore filling can occur in series or in parallel and depend on electrolyte composition, morphology, and state of charge [7–10]. Accordingly, no transfer coefficient from a solution-phase standard electron-transfer rate constant to charge-transfer resistance, chemical diffusion, capacity, rate capability, or cycle life is introduced here.",
    )
    ensure_before(
        doc,
        "2.2.",
        "Future tests should pre-specify endpoint-specific",
        "Future tests should pre-specify endpoint-specific effect sizes rather than a universal multiplier. Examples include ln(Rct,control/Rct,tBLG), ln(DNa,tBLG/DNa,control), the first-cycle Coulombic-efficiency difference, and the capacity-retention difference at fixed loading and protocol. Minimum important differences should be set from measurement precision, statistical power, and application requirements, with confidence intervals across independent material batches. Controls should include uncoated hard carbon, an equal-carbon-loading conventional coating, AB or untwisted bilayer graphene, and the proposed tBLG coating.",
    )
    replace_matching(doc, ("2.2. Screening-level cost", "2.2. Illustrative parametric"), "2.2. Illustrative parametric cost stress test", "Heading 3")
    replace_matching(
        doc,
        ("The functional unit is",),
        "The functional unit is 1 kWh of nominal battery-system capacity at the factory gate. The archived accounting envelope contains raw materials, a hypothetical CVD growth-and-transfer coating category, cell assembly, balance of system (BOS), and overhead. Installation, financing, operation, maintenance, replacement, and end-of-life treatment are outside the boundary. The output is an illustrative parametric stress test in USD/kWh, not a vendor quote, bottom-up plant TEA, net-present-value model, levelized cost of storage, or competitiveness forecast.",
    )
    replace_matching(
        doc,
        ("The boundary includes",),
        "Because no tBLG-on-powder unit operation is defined in the archive or reviewed sources, the model contains no process-specific mass or energy balance, equipment sizing, location, operating hours, throughput, yield loss, direct CapEx, direct OpEx, recycle loop, or quality-control burden. It also assigns no economic credit for power, lifetime, thermal management, utilization, or replacement because Levels 3 and 4 are unmet. These omissions are decision-critical data gaps and prevent classification as a process TEA [28].",
    )
    replace_matching(
        doc,
        ("Before the multiplier, raw materials",),
        "Before the multiplier, raw materials, cell assembly, BOS, and overhead are sampled from normal distributions with means of 40, 25, 20, and 15 USD/kWh and standard deviations of 8, 5, 4, and 3 USD/kWh, respectively. CVD coating is sampled from a lognormal distribution with a median of 35 USD/kWh and log-space sigma of 0.4. These are archived screening assumptions, not independently validated process estimates. Table A.1 reports the exact distributions and evidence status.",
    )
    replace_matching(doc, ("2.3. Technology innovation systems analysis", "2.3. Innovation-system"), "2.3. Innovation-system evidence and provenance audit", "Heading 3")
    replace_matching(
        doc,
        ("The seven TIS functions defined",),
        "The TIS component was converted from a numerical diagnosis to a provenance audit. The function definitions are available from the published framework [23,24], and battery-sector studies illustrate the empirical records normally required [25,26]. The project archive, however, retains only aggregate author-coded labels and short narratives; it does not contain the underlying publication, patent, firm, roadmap, or standards records used to construct them.",
    )
    replace_matching(
        doc,
        ("For interpretation, 1 denotes",),
        "The audit records whether each required evidence item is available, not archived, or not collected. Items include a reproducible search strategy and inclusion criteria, raw indicator records, a codebook, weights and normalization, independent coding and inter-rater assessment, and stakeholder or expert elicitation. Table A.2 provides the resulting evidence-availability matrix.",
    )
    replace_matching(
        doc,
        ("No statistical confidence interval",),
        "Legacy TIS_Scores.csv and Monte_Carlo_TIS_10000.csv are retained in the supplementary package solely to document the earlier workflow. Resampling author-assigned scores cannot create empirical uncertainty, so neither file is used to rank functions, technologies, countries, or policy options.",
        "Normal",
    )
    replace_matching(
        doc,
        ("The archived TIS Monte Carlo file",),
        "The only permitted inference is about evidence readiness: the available archive cannot support a quantitative commercialization assessment or policy recommendation. Such inference would require transparent indicators, reproducible coding, uncertainty analysis tied to observed data, and independent stakeholder input.",
    )

    replace_matching(
        doc,
        ("Reviewer-identified evidence gaps",),
        "Reviewer-identified gaps were organized into sequential decision gates: relevant non-aqueous electrochemistry, particle-level manufacturability, controlled half-cell validation, practical full-cell value, process and life-cycle inventory, and innovation-system evidence. A gate is passed only when pre-specified measurements and provenance records are available. No funding amount, institutional lead, standards schedule, procurement instrument, or other policy action is recommended.",
    )

    replace_matching(doc, ("3.1. Interfacial enhancement scenarios", "3.1. Model-system"), "3.1. Model-system observation and endpoint-specific validation", "Heading 3")
    replace_matching(
        doc,
        ("Because the reviewed sources",),
        "The reported twist dependence establishes that the electronic structure of a well-defined planar tBLG electrode can modulate weak-coupling outer-sphere electron transfer under the studied aqueous conditions [18,19]. It does not establish the direction or magnitude of any hard-carbon SIB endpoint. Fig. 1 therefore delineates the boundary between the supported planar model system and untested particle, cell, process, commercial, and policy domains. The proposed coating remains at Level 0 until relevant non-aqueous and particle-level evidence is obtained.",
    )
    # Rebuild the Figure 1 block explicitly.  The original picture paragraph
    # has no text; treating it as a blank text paragraph deletes the drawing.
    fig1_caption = find_paragraph(doc, "Fig. 1.")
    previous_xml = fig1_caption._p.getprevious()
    while isinstance(previous_xml, CT_P):
        previous = Paragraph(previous_xml, fig1_caption._parent)
        is_old_block = previous.text.startswith("A future half-cell study should report")
        is_picture = bool(previous._p.xpath(".//w:drawing"))
        if not (is_old_block or is_picture):
            break
        previous_xml = previous_xml.getprevious()
        previous._p.getparent().remove(previous._p)

    future = fig1_caption.insert_paragraph_before(
        "A future half-cell study should report endpoint-specific effect sizes and confidence intervals across independent batches. Failure to exceed a pre-specified minimum important difference, or evidence of pore blockage, unstable twist structure, or adverse SEI growth, would falsify the progression hypothesis. A positive half-cell result would justify a practical full-cell test, not a commercial or policy conclusion.",
        "Normal",
    )
    future.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    future.paragraph_format.first_line_indent = Inches(0.33)
    future.paragraph_format.line_spacing = 1.5
    future.paragraph_format.keep_with_next = False

    picture = fig1_caption.insert_paragraph_before()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.space_before = Pt(3)
    picture.paragraph_format.space_after = Pt(3)
    picture.paragraph_format.keep_with_next = True
    picture.add_run().add_picture(
        str(SUBMISSION_FIGURES / "Figure_1.png"),
        width=Inches(6.25),
    )
    set_text(
        fig1_caption,
        "Fig. 1. Evidence boundary for the proposed tBLG-coated hard-carbon SIB concept. Left: the supported planar aqueous outer-sphere model system. Center and right: untested particle, cell, process, commercial, and policy domains. The illustration is not a development sequence or readiness forecast.",
        "Caption",
    )
    replace_matching(
        doc,
        ("Table 1 presents",),
        "Table 1 states the seven evidence levels, their current status, and the inference permitted at each level. It replaces the earlier cross-source performance and cost comparison, whose active-material, cell, and system boundaries were not harmonized.",
    )
    replace_matching(doc, ("Illustrative cross-source benchmarks", "Evidence hierarchy and permitted"), "Evidence hierarchy and permitted inference for the proposed tBLG-coated hard-carbon SIB concept.", "Caption")

    replace_matching(doc, ("3.2. Screening cost envelope", "3.2. Parametric cost"), "3.2. Parametric cost stress test and input-cost associations", "Heading 3")
    replace_matching(
        doc,
        ("The archived trajectory file assigns",),
        "The archived trajectory file contains deterministic 2025–2035 cost inputs for a conventional hard-carbon SIB reference and the hypothetical tBLG scenario. They are retained as a provenance view of the earlier assumptions, not as harmonized technology forecasts or a battery-competitiveness comparison. Fig. 2 pairs those inputs with a reproducible Spearman rank-correlation screen of the separate 10,000-run 2035 Monte Carlo dataset.",
    )
    replace_matching(
        doc,
        ("Fig. 2.",),
        "Fig. 2. Archived cost trajectories and input-cost rank associations. (a) Deterministic scenario inputs, not statistical forecasts. (b) Spearman rank correlations calculated from the 10,000 archived 2035 runs.",
        "Caption",
    )
    replace_matching(
        doc,
        ("The 2035 screening distribution has",),
        "The 2035 stress-test distribution has a mean of 121.4 USD/kWh, median of 119.5 USD/kWh, standard deviation of 17.4 USD/kWh, and P5–P95 range of 96.5–152.6 USD/kWh (Fig. 3). After applying the learning multiplier, the mean component costs are 35.18 USD/kWh for raw materials, 33.43 for CVD coating, 21.97 for cell assembly, 17.63 for BOS, and 13.21 for overhead. The mean run-level shares of raw materials and CVD coating sum to 56.0%. Table 2 reports component means, P5–P95 ranges, and mean run-level shares.",
    )
    replace_matching(
        doc,
        ("Fig. 3.",),
        "Fig. 3. Monte Carlo screening results from 10,000 archived runs. (a) Total 2035 scenario-cost distribution with mean, median, P5, and P95 markers. (b) Learning-adjusted component means with P5–P95 intervals; CVD coating has the broadest component interval.",
        "Caption",
    )

    replace_matching(doc, ("3.4. Conditional cost-ratio illustration", "3.4. Separation"), "3.4. Separation of model-system and economic evidence", "Heading 3")
    replace_matching(
        doc,
        ("To expose the scale", "The model-system observation and"),
        "The model-system observation and the archived cost stress test answer different questions and are deliberately not coupled. No kinetic, capacity, rate, or lifetime credit is entered into the cost arithmetic, and the non-harmonized tBLG, conventional SIB, and LFP values are not divided to produce break-even ratios. A future application model would require measured full-cell power, efficiency, lifetime, thermal, utilization, and replacement effects under a common functional unit and system boundary. Until Levels 3–5 are met, neither a cost premium nor competitiveness can be inferred from the planar aqueous result.",
    )
    replace_matching(
        doc,
        ("Fig. 4.",),
        "Fig. 4. Independent evidence domains for an integrated assessment of the proposed tBLG-SIB concept. The central unvalidated material concept is surrounded by six domains that require direct measurements or traceable records. Dashed links are non-directional interfaces, not transfer coefficients, process steps, or readiness levels.",
        "Caption",
    )
    replace_matching(
        doc,
        ("The reviewed sources and project files provide no demonstration",),
        "The reviewed sources and project files provide no demonstration of a continuous, near-magic-angle bilayer on irregular hard-carbon particles. Existing studies show twist control and CVD growth on comparatively flat substrates [15,21,22], not conformal powder coverage. Critical measurements include coverage, local twist distribution, defects, adhesion, transfer residue, powder and slurry compatibility, calendaring survival, wetting, pore accessibility, SEI chemistry, Na+ transport, and cycling stability. Until a workable unit operation and acceptable yield are demonstrated, plant scale, energy use, material balance, CapEx, and OpEx cannot be estimated defensibly.",
    )

    replace_matching(doc, ("4. Exploratory TIS diagnosis", "4. Innovation-system evidence audit"), "4. Innovation-system evidence audit and evidence-gated development", "Heading 2")
    replace_matching(doc, ("4.1. TIS functional assessment", "4.1. Provenance audit"), "4.1. Provenance audit", "Heading 3")
    replace_matching(
        doc,
        ("The archived author-coded scores",),
        "The provenance audit found that the archive contains the TIS function labels and legacy aggregate scores but not the raw observations, search protocol, codebook, weights, normalization, independent coding, or stakeholder evidence needed to reproduce or interpret those scores. The decimal values are therefore removed from the manuscript's results.",
    )
    replace_matching(
        doc,
        ("The scores are ordinal",),
        "Fig. 5 visualizes the archive composition, while Table A.2 enumerates evidence availability. Together they show which records are present and which must be collected before an innovation-system comparison is attempted; they do not estimate uncertainty, readiness, market formation, or commercial viability.",
    )
    replace_matching(
        doc,
        ("Fig. 5.",),
        "Fig. 5. Visual provenance map of the archived innovation-system materials. The archive contains TIS definitions and legacy aggregate labels, whereas raw indicators, search protocol, codebook and weights, independent coding, and stakeholder evidence are missing. The supported result is limited to provenance gaps; no ranking, commercialization claim, or policy recommendation is permitted.",
        "Caption",
    )
    replace_matching(doc, ("4.2. Interpretation and limitations", "4.2. Permitted use"), "4.2. Permitted use of the innovation-system framework", "Heading 3")
    replace_matching(
        doc,
        ("The TIS framework is used",),
        "The TIS framework is used only to organize data-acquisition questions. A system-level observation cannot validate an electrochemical premise, and missing innovation-system records cannot be converted into measured weaknesses. The audit therefore supports no technology or country ranking and no statement that a particular TIS function is strong or weak.",
    )
    replace_matching(
        doc,
        ("Because the archive lacks",),
        "Because the archive lacks raw indicators and stakeholder evidence, the study cannot assign leadership, funding, standards timelines, market incentives, or procurement actions. If the technical concept reaches later gates, a future study should publish its search protocol and indicator records, predefine the codebook and weights, use independent coders, report reliability, and conduct structured stakeholder or expert elicitation before drawing commercial or policy conclusions.",
    )
    replace_matching(doc, ("4.3. Milestone-gated", "4.3. Evidence-gated"), "4.3. Evidence-gated research agenda", "Heading 3")
    replace_matching(
        doc,
        ("Gate 1 concerns",),
        "Gate 1 concerns the physical premise. Twist-dependent electron transfer must first be reproduced in a relevant non-aqueous electrolyte model system with matched AB-stacked, untwisted, and conventional-carbon controls; the aqueous outer-sphere observation alone does not pass this gate [18,19].",
    )
    replace_matching(
        doc,
        ("Gate 2 concerns",),
        "Gate 2 concerns particle-level manufacturability. Researchers must demonstrate reproducible bilayer coverage on hard-carbon particles, quantify twist-angle and defect distributions, verify adhesion through mixing and calendaring, and show that the coating does not occlude the pore structure [15,21,22].",
    )
    replace_matching(
        doc,
        ("Gate 3 concerns",),
        "Gate 3 concerns half-cell electrochemical validity. Independent material batches must compare uncoated hard carbon, an equal-carbon-loading conventional coating, AB or untwisted bilayer graphene, and tBLG under a shared non-aqueous protocol. Primary endpoint effect sizes and confidence intervals should cover impedance, diffusion, Coulombic efficiency, capacity retention, rate response, degradation, and SEI chemistry [7–10,20,30]. Failure to exceed the pre-specified minimum important difference would not support progression.",
    )
    replace_matching(
        doc,
        ("Gate 4 concerns",),
        "Gate 4 concerns practical full-cell value. Only after a reproducible half-cell benefit should full cells be tested at practical mass loading, N/P ratio, electrode density, and electrolyte quantity for power, cycle and calendar life, efficiency, thermal behavior, and safety under a defined application profile [30].",
    )
    replace_matching(
        doc,
        ("Gate 5 concerns",),
        "Gate 5 concerns process economics, life-cycle inventory, and application-specific system value. A defined coating route must provide material and energy balances, yield and recycle assumptions, throughput, operating hours, location, equipment sizing, CapEx, OpEx, waste streams, and quality control before bottom-up TEA or LCA is attempted; validated full-cell data are also required for an application model [2,27–29,31,32].",
    )

    replace_matching(
        doc,
        ("The principal limitation is",),
        "The principal limitation is that no tBLG-coated hard-carbon SIB fabrication or test result was identified in the reviewed sources or project files. The aqueous solution-phase response reported for planar tBLG [18,19] cannot be equated with Na+ desolvation, SEI transport, adsorption, insertion, pore filling, or degradation. The earlier universal kinetic multipliers are removed rather than treated as modeled results.",
    )
    replace_matching(
        doc,
        ("Manufacturability is equally uncertain",),
        "Manufacturability is equally uncertain. Controlled tBLG is documented primarily on comparatively flat substrates [15,21,22], not as a conformal, twist-controlled coating on battery powder. The effects of curvature, defects, transfer contamination, SEI coverage, pore blocking, slurry processing, and cycling stress remain unknown.",
    )
    replace_matching(
        doc,
        ("The cost analysis is a parametric",),
        "The cost analysis is an illustrative parametric stress test built from hard-coded accounting categories and a bounded learning multiplier. It is not linked to a process flow sheet, plant design, material or energy balance, coating yield, equipment quote, or direct CapEx/OpEx model. The input distributions are subjective scenario ranges, so the 2035 distribution is conditional rather than predictive.",
    )
    replace_matching(
        doc,
        ("The TIS analysis is based",),
        "The innovation-system portion is limited to a provenance audit. Legacy aggregate scores remain in the supplementary archive for transparency, but the missing raw records, normalization, weights, coder agreement, and stakeholder elicitation prevent quantitative TIS inference.",
    )
    replace_matching(
        doc,
        ("No life-cycle assessment was conducted",),
        "No life-cycle assessment was conducted. Energy use, gases, substrates, transfer chemicals, yield, and waste for a scalable coating route are unknown. Published battery LCA studies provide methodological baselines [31,32] but cannot substitute for a tBLG-powder process inventory.",
    )
    replace_matching(
        doc,
        ("Finally, the break-even calculation", "Finally, no bridge"),
        "Finally, no bridge is estimated between the Level 0 physical premise and the cost stress test. The manuscript therefore does not quantify the economic value of a kinetic signal, a break-even lifetime, commercial competitiveness, or a policy benefit. These questions become identifiable only after validated full-cell data and a harmonized application model exist.",
    )

    replace_matching(
        doc,
        ("This study is an early-stage",),
        "This study is an evidence-gated assessment of a proposed and unvalidated tBLG coating for hard-carbon SIB anodes. It does not demonstrate a coating, a sodium-storage mechanism, improved cell performance, a scalable process, commercial competitiveness, or policy readiness.",
    )
    replace_matching(
        doc,
        ("Conditional on the archived",),
        "Conditional on the archived cost assumptions, the 2035 stress-test distribution has a median of 119.5 USD/kWh and a P5–P95 range of 96.5–152.6 USD/kWh. Raw materials and CVD coating have mean run-level shares totaling 56.0%, and CVD cost has the largest absolute Spearman rank correlation with total scenario cost. These are properties of the assumed inputs and formula, not process-derived forecasts or evidence of real cost drivers.",
    )
    replace_matching(
        doc,
        ("The archived author-coded TIS rubric",),
        "The innovation-system provenance audit shows that the records needed to reproduce numerical TIS scores or support policy inference are unavailable. The legacy scores are therefore excluded from the results; the framework is retained only to define future evidence requirements.",
    )
    replace_matching(
        doc,
        ("The next decision should be experimental",),
        "The next decision should be experimental. Relevant non-aqueous model-electrode tests and particle-level coating verification precede any hard-carbon half-cell claim. Endpoint-specific effect sizes with confidence intervals and matched controls should determine progression; a positive half-cell result would justify practical full-cell testing, not immediate commercialization.",
    )
    replace_matching(
        doc,
        ("Only after full-cell performance",),
        "Only after full-cell performance and a defined, reproducible coating process are available can material and energy balances, CapEx, OpEx, life-cycle impacts, system-level value, commercial evidence, and stakeholder-based innovation-system analysis be integrated into a defensible technology assessment.",
    )

    replace_matching(
        doc,
        ("Qinyou Yang: Writing",),
        "Qinyou Yang: Writing – original draft, Writing – review & editing, Conceptualization, Methodology, Investigation, Software, Formal analysis, Data curation, Resources, Visualization, Supervision. Xinran Wang: Investigation, Validation. Niwei Zhu: Validation. Wulin Zhang: Validation, Funding acquisition.",
    )
    replace_matching(doc, ("Declaration of competing interests", "Declaration of competing interest"), "Declaration of competing interest", "Heading 2")
    replace_matching(doc, ("Acknowledgements", "Funding"), "Funding", "Heading 2")
    replace_matching(
        doc,
        ("The authors gratefully acknowledge", "This work was supported"),
        "This work was supported by the Key Project of Heilongjiang Provincial Natural Science Foundation United Fund (Grant No. ZL2025C004). The funder had no role in study design, data analysis, interpretation, manuscript preparation, or the decision to submit.",
    )
    # Elsevier asks that the AI declaration be immediately before the references.
    data_heading = find_paragraph(doc, "Data availability")
    data_body = next_paragraph(data_heading)
    ai_heading = find_paragraph(
        doc,
        "AI use declaration",
        "Declaration of generative AI and AI-assisted technologies",
    )
    ai_body = next_paragraph(ai_heading)
    set_text(data_heading, "Data availability", "Heading 2")
    set_text(
        data_body,
        "The screening inputs, deterministic source code, recomputed summaries, and evidence-provenance table used in this revision are provided with the submission as Supplementary Material. They are modeled inputs and outputs, not measured tBLG-SIB data. Zenodo v1.0.1 [33] is the currently public version under the unchanged concept DOI:",
        "Normal",
    )
    data_body.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doi_paragraph = insert_paragraph_before(
        ai_heading,
        "https://doi.org/10.5281/zenodo.19362805",
        "Normal",
    )
    data_continuation = insert_paragraph_before(
        ai_heading,
        "This public version predates the present revision and must not be represented as containing the revised supplementary package. The revised package is prepared for release as a new version under the same concept DOI before resubmission.",
        "Normal",
    )
    for paragraph in (doi_paragraph, data_continuation):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Inches(0)
    doi_run = doi_paragraph.runs[0]
    doi_run.font.name = "Times New Roman"
    doi_run.font.size = Pt(12)
    doi_fonts = doi_run._element.get_or_add_rPr().get_or_add_rFonts()
    doi_fonts.set(qn("w:ascii"), "Times New Roman")
    doi_fonts.set(qn("w:hAnsi"), "Times New Roman")
    set_text(
        ai_heading,
        "Declaration of generative AI and AI-assisted technologies in the manuscript preparation process",
        "Heading 2",
    )
    set_text(
        ai_body,
        "During the preparation of this work, the authors used Consensus to support literature discovery and OpenAI Codex to support manuscript organization, reference cross-checking, code review, and language/readability revision. No AI tool was used to create experimental data. All reported numerical outputs were generated by the archived deterministic code and checked against the source files. After using these tools, the authors reviewed and edited the content as needed and take full responsibility for the content of the published article.",
        "Normal",
    )
    ai_body.alignment = WD_ALIGN_PARAGRAPH.LEFT
    references_heading = find_paragraph(doc, "References")
    for paragraph in (
        data_heading,
        data_body,
        doi_paragraph,
        data_continuation,
        ai_heading,
        ai_body,
    ):
        references_heading._p.addprevious(paragraph._p)

    # Remove the obsolete TIS score symbol from the nomenclature table.
    for row in list(doc.tables[0].rows):
        if row.cells[0].text.strip() == "Sj":
            remove_row(row)
    for row in doc.tables[0].rows:
        if row.cells[0].text.strip() in {"K0", "k0", "k₀"}:
            row.cells[0].text = "k₀"

    set_table_data(
        doc.tables[1],
        [
            ["Level", "Evidence required", "Current status", "Permitted inference"],
            ["0", "Planar tBLG response of aqueous solution-phase outer-sphere probes", "Supported for the published model system [18,19]", "Twist-sensitive heterogeneous electron transfer under the reported model conditions only"],
            ["1", "Replication in a relevant non-aqueous electrolyte model system", "No direct evidence identified", "Experimental hypothesis only"],
            ["2", "Particle coating, coverage, twist, adhesion, pore access, and structural stability", "tBLG fabrication is reported on comparatively flat substrates [15,21,22], not hard-carbon powder", "No manufacturability claim"],
            ["3", "Matched hard-carbon half-cell endpoints across independent material batches", "Hard-carbon interface methods exist [7–10,20], but no tBLG-coated hard-carbon test was identified", "No Na-storage, capacity, rate, impedance, SEI, or lifetime claim"],
            ["4", "Practical full-cell performance and safety under a defined application profile", "Relevant SIB pouch-cell methods exist [30], but no tBLG full cell was identified", "No power, energy, efficiency, safety, or system-life claim"],
            ["5", "Process inventory, plant model, LCA boundary, and application-linked value model", "General battery cost/LCA methods exist [2,27–29,31,32]; no tBLG-powder process exists", "Archived parametric cost stress test only; no process TEA, LCA, or competitiveness claim"],
            ["6", "Traceable innovation-system indicators, reproducible coding, and stakeholder evidence", "Framework and battery precedents exist [23–26]; concept-specific records are incomplete", "Provenance audit only; no commercial ranking or policy recommendation"],
        ],
    )
    set_table_data(
        doc.tables[4],
        [
            ["Evidence item", "Evidence status", "Consequence", "Permitted use"],
            ["TIS function definitions", "Available from published literature [23,24]", "Analytical questions can be defined", "Conceptual framework only"],
            ["Aggregate author-coded labels", "Available in legacy TIS_Scores.csv", "Score construction and decimal precision are unsupported", "Provenance record only"],
            ["Raw indicator records", "Not archived", "Claims cannot be independently reproduced", "None"],
            ["Search strategy and inclusion criteria", "Not archived", "Selection cannot be audited", "None"],
            ["Codebook, weights, and normalization", "Not archived", "Scores cannot be reconstructed", "None"],
            ["Independent coding and reliability", "Not conducted", "Coder dependence is unknown", "None"],
            ["Stakeholder or expert elicitation", "Not conducted", "No empirical basis for commercial or policy inference", "None"],
        ],
    )

    # Typographic consistency for ranges in the reported and archived tables.
    for row in doc.tables[2].rows[1:]:
        row.cells[2].text = row.cells[2].text.replace("-", "–")
    for row in doc.tables[3].rows[1:]:
        row.cells[2].text = row.cells[2].text.replace("1-100", "1–100").replace("10%-20%", "10%–20%")

    replace_matching(doc, ("Table A ", "Table A.1"), "Table A.1", "Caption")
    replace_matching(doc, ("Archived Monte Carlo cost-input",), "Archived Monte Carlo cost-input distributions and evidence status.", "Caption")
    replace_matching(doc, ("Table B ", "Table A.2"), "Table A.2", "Caption")
    replace_matching(doc, ("Exploratory TIS scores", "Innovation-system evidence"), "Innovation-system evidence-availability and provenance audit.", "Caption")

    set_reference_list(doc, REFERENCES)
    mapping = renumber_citations_and_references(doc)
    for table in doc.tables:
        as_three_line_table(table)
        repeat_header(table.rows[0])

    # Compact, readable layouts prevent orphaned captions, mid-word wrapping,
    # and single-row continuation pages in Word/LibreOffice rendering.
    repeat_header(doc.tables[0].rows[1])
    set_column_widths(doc.tables[0], [2.15, 4.45])
    format_table_text(doc.tables[0], 9.0)
    for row in doc.tables[0].rows:
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    insert_page_break_before_table(doc.tables[0])

    set_column_widths(doc.tables[1], [0.48, 1.90, 2.30, 1.92])
    format_table_text(doc.tables[1], 9.0)
    set_column_widths(doc.tables[4], [1.55, 1.80, 2.25, 1.10])
    format_table_text(doc.tables[4], 8.5)

    format_reference_list(doc)
    remove_last_rendered_page_breaks(doc)
    add_continuous_line_numbers(doc)
    doc.save(MANUSCRIPT)
    return mapping


def revise_highlights() -> None:
    highlights = [
        "Planar aqueous tBLG redox kinetics are not mapped to sodium-ion cells.",
        "Only the planar model-system premise currently has direct evidence.",
        "A 10,000-run cost stress test propagates assumptions, not forecasts.",
        "Six evidence gates precede process TEA, commercialization, or policy claims.",
    ]
    if not 3 <= len(highlights) <= 5 or any(len(item) > 85 for item in highlights):
        raise ValueError("Highlights must contain 3–5 items of at most 85 characters.")
    doc = Document(HIGHLIGHTS)
    while len(doc.paragraphs) < len(highlights):
        doc.add_paragraph()
    for paragraph, item in zip(doc.paragraphs, highlights):
        set_text(paragraph, item)
        paragraph.style = "List Bullet"
    for paragraph in doc.paragraphs[len(highlights) :]:
        paragraph._element.getparent().remove(paragraph._element)
    doc.save(HIGHLIGHTS)


def main() -> None:
    mapping = revise_manuscript()
    revise_highlights()
    print(f"Revised manuscript and highlights. Citation mapping: {mapping}")


if __name__ == "__main__":
    main()
