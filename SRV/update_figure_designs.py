#!/usr/bin/env python3
"""Update figure-related prose after the non-flowchart artwork redesign."""

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
MANUSCRIPT = ROOT / "Manuscript.docx"
RESPONSE = ROOT / "Response_to_Reviewers.docx"


MANUSCRIPT_UPDATES = {
    "The reported twist dependence establishes": (
        "The reported twist dependence establishes that the electronic structure of a well-defined planar tBLG electrode can modulate weak-coupling outer-sphere electron transfer under the studied aqueous conditions [18,19]. It does not establish the direction or magnitude of any hard-carbon SIB endpoint. Fig. 1 therefore delineates the boundary between the supported planar model system and untested particle, cell, process, commercial, and policy domains. The proposed coating remains at Level 0 until relevant non-aqueous and particle-level evidence is obtained."
    ),
    "Fig. 1.": (
        "Fig. 1. Evidence boundary for the proposed tBLG-coated hard-carbon SIB concept. Left: the supported planar aqueous outer-sphere model system. Center and right: untested particle, cell, process, commercial, and policy domains. The illustration is not a development sequence or readiness forecast."
    ),
    "Fig. 2.": (
        "Fig. 2. Archived cost trajectories and input-cost rank associations. (a) Deterministic scenario inputs, not statistical forecasts. (b) Spearman rank correlations calculated from the 10,000 archived 2035 runs."
    ),
    "Fig. 3.": (
        "Fig. 3. Monte Carlo screening results from 10,000 archived runs. (a) Total 2035 scenario-cost distribution with mean, median, P5, and P95 markers. (b) Learning-adjusted component means with P5–P95 intervals; CVD coating has the broadest component interval."
    ),
    "Fig. 4.": (
        "Fig. 4. Independent evidence domains for an integrated assessment of the proposed tBLG-SIB concept. The central unvalidated material concept is surrounded by six domains that require direct measurements or traceable records. Dashed links are non-directional interfaces, not transfer coefficients, process steps, or readiness levels."
    ),
    "Fig. 5 and Table A.2": (
        "Fig. 5 visualizes the archive composition, while Table A.2 enumerates evidence availability. Together they show which records are present and which must be collected before an innovation-system comparison is attempted; they do not estimate uncertainty, readiness, market formation, or commercial viability."
    ),
    "Fig. 5.": (
        "Fig. 5. Visual provenance map of the archived innovation-system materials. The archive contains TIS definitions and legacy aggregate labels, whereas raw indicators, search protocol, codebook and weights, independent coding, and stakeholder evidence are missing. The supported result is limited to provenance gaps; no ranking, commercialization claim, or policy recommendation is permitted."
    ),
}

RESPONSE_UPDATES = {
    "Response: Agreed. Figure 1 contains no dates": (
        "Response: Agreed. Figure 1 contains no dates, TRLs, or statistical bands. It now uses a non-sequential evidence-boundary schematic to separate the supported planar model system from untested particle, cell, process, commercial, and policy domains."
    )
}


def apply_updates(path: Path, updates: dict[str, str]) -> None:
    document = Document(path)
    found = set()
    for paragraph in document.paragraphs:
        for prefix, replacement in updates.items():
            if paragraph.text.startswith(prefix):
                paragraph.clear()
                paragraph.add_run(replacement)
                found.add(prefix)
                break
    missing = set(updates) - found
    if missing:
        raise ValueError(f"Missing expected paragraphs in {path.name}: {sorted(missing)}")
    document.save(path)


def main() -> None:
    apply_updates(MANUSCRIPT, MANUSCRIPT_UPDATES)
    apply_updates(RESPONSE, RESPONSE_UPDATES)
    print("Updated manuscript captions, figure discussion, and reviewer response.")


if __name__ == "__main__":
    main()
